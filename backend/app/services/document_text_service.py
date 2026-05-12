"""Extração de texto de documentos para alimentar a análise consolidada.

Estratégia por tipo de arquivo:

* PDF        : ``pypdf`` primeiro (texto digital). Se o resultado for
               pobre (menos de ~30 chars/página ou < 200 chars totais),
               o caller cai para Vision (rendering via ``pypdfium2``).
* JPG/PNG    : sem texto digital — vão direto como imagem para o LLM.
* DOCX       : ``python-docx`` (parágrafos + tabelas concatenados).
* TXT        : ``decode('utf-8')`` com fallback latin-1.

Tudo são funções puras. Não levantam para erros conhecidos — retornam
string vazia e deixam o caller decidir o que fazer (ex.: marcar `notes`
no report).
"""

from __future__ import annotations

import io
from typing import Final

import pypdfium2 as pdfium  # type: ignore[import-untyped]

from app.core.logging import get_logger

logger = get_logger(__name__)

# Limite duro para evitar que um único PDF/DOCX gigante explode o context
# do LLM. 200k chars ≈ 50k tokens — já é o teto operacional do gpt-4o.
MAX_TEXT_CHARS: Final[int] = 200_000

# Heurísticas para decidir "este PDF tem texto digital extraível?"
MIN_CHARS_PER_PAGE: Final[int] = 30
MIN_TOTAL_CHARS: Final[int] = 200

# Render de PDF → PNG (reaproveita parâmetros do matricula_ocr_service).
RENDER_SCALE: Final[float] = 200.0 / 72.0
MAX_PAGES_RENDER: Final[int] = 8


def extract_text_from_pdf(pdf_bytes: bytes) -> tuple[str, int]:
    """Tenta extrair texto digital de um PDF via ``pypdf``.

    Retorna ``(texto, n_paginas)``. Em qualquer falha, devolve
    ``("", 0)`` — caller decide se chama Vision OCR.
    """
    if not pdf_bytes:
        return "", 0
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("document_text.pypdf_missing")
        return "", 0
    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    except Exception as exc:  # noqa: BLE001
        logger.warning("document_text.pdf_open_failed", error=str(exc))
        return "", 0

    parts: list[str] = []
    n = len(reader.pages)
    try:
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception as exc:  # noqa: BLE001
                logger.warning("document_text.pdf_page_failed", error=str(exc))
                parts.append("")
    finally:
        # PdfReader não precisa close() explícito.
        pass

    text = "\n\n".join(p.strip() for p in parts if p and p.strip())
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[... truncado ...]"
    return text, n


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extrai texto de um .docx (Open XML) — parágrafos + tabelas.

    Falha silenciosa: arquivo corrompido ou .doc binário (pré-2007) →
    string vazia. Caller pode validar via MIME antes de chamar.
    """
    if not docx_bytes:
        return ""
    try:
        from docx import Document  # type: ignore[import-untyped]
    except ImportError:
        logger.warning("document_text.python_docx_missing")
        return ""
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as exc:  # noqa: BLE001
        logger.warning("document_text.docx_open_failed", error=str(exc))
        return ""

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts)
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[... truncado ...]"
    return text


def extract_text_from_txt(txt_bytes: bytes) -> str:
    """Decode UTF-8 com fallback latin-1. Trunca em ``MAX_TEXT_CHARS``."""
    if not txt_bytes:
        return ""
    try:
        text = txt_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = txt_bytes.decode("latin-1")
        except UnicodeDecodeError:
            text = txt_bytes.decode("utf-8", errors="replace")
    if len(text) > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS] + "\n\n[... truncado ...]"
    return text


def needs_vision_ocr(text: str, page_count: int) -> bool:
    """``True`` quando o texto extraído via ``pypdf`` é pobre demais e
    o caller deve cair para Vision OCR.
    """
    if page_count <= 0:
        return False
    if not text:
        return True
    if len(text) < MIN_TOTAL_CHARS:
        return True
    chars_per_page = len(text) / max(page_count, 1)
    return chars_per_page < MIN_CHARS_PER_PAGE


def render_pdf_to_pngs(
    pdf_bytes: bytes, *, max_pages: int = MAX_PAGES_RENDER
) -> list[bytes]:
    """Renderiza até ``max_pages`` páginas em PNG.

    Reaproveita o mesmo motor (``pypdfium2``) do ``matricula_ocr_service``.
    Em qualquer falha devolve lista vazia — caller registra `notes`.
    """
    if not pdf_bytes:
        return []
    try:
        doc = pdfium.PdfDocument(pdf_bytes)
    except Exception as exc:  # noqa: BLE001
        logger.warning("document_text.render_open_failed", error=str(exc))
        return []

    pages: list[bytes] = []
    n = min(len(doc), max_pages)
    try:
        for i in range(n):
            page = doc[i]
            try:
                pil = page.render(scale=RENDER_SCALE).to_pil()
            finally:
                page.close()
            buf = io.BytesIO()
            pil.save(buf, format="PNG", optimize=True)
            pages.append(buf.getvalue())
    finally:
        doc.close()
    return pages


__all__ = [
    "MAX_PAGES_RENDER",
    "MAX_TEXT_CHARS",
    "MIN_CHARS_PER_PAGE",
    "MIN_TOTAL_CHARS",
    "extract_text_from_docx",
    "extract_text_from_pdf",
    "extract_text_from_txt",
    "needs_vision_ocr",
    "render_pdf_to_pngs",
]
